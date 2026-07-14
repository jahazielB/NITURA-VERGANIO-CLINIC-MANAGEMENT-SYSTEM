from __future__ import annotations

import math
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import legal
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, Table, TableStyle
from reportlab.lib.units import inch

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output"
DOCX_PATH = OUT_DIR / "PSA-CRASM-Endorsement-United-Fresh-Fire-Church.docx"
PDF_PATH = OUT_DIR / "PSA-CRASM-Endorsement-United-Fresh-Fire-Church.pdf"
LOGO_PATH = ROOT / "src" / "assets" / "NV-logo.png"
ORANGE = "#F58220"
ORANGE_RGB = colors.HexColor(ORANGE)


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("left", "top", "right", "bottom", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            element = tcBorders.find(qn(f"w:{edge}"))
            if element is None:
                element = OxmlElement(f"w:{edge}")
                tcBorders.append(element)
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "FFFFFF")


def set_row_height(row, height_inches):
    row.height = Inches(height_inches)
    row.height_rule = WD_SECTION_START.NEW_PAGE if False else None


def add_paragraph(cell_or_doc, text="", *, bold=False, color=None, size=11, align=WD_ALIGN_PARAGRAPH.LEFT,
                  spacing_before=0, spacing_after=0, line_spacing=1.0, underline=False, italic=False):
    p = cell_or_doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(spacing_before)
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.line_spacing = line_spacing
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        if color:
            r.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
        r.font.size = Pt(size)
        r.font.name = "Calibri"
        r.underline = underline
    return p


def add_run(paragraph, text="", *, bold=False, color=None, size=11, underline=False, italic=False):
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor.from_string(color.replace("#", ""))
    r.font.size = Pt(size)
    r.font.name = "Calibri"
    r.underline = underline
    return r


def make_docx() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(13)
    for margin_name in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, margin_name, Inches(0.75))

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    available_width = 7.0
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(1.35)
    table.columns[1].width = Inches(5.65)
    remove_table_borders(table)
    logo_cell, text_cell = table.rows[0].cells
    set_cell_margins(logo_cell, 0, 0, 0, 0)
    set_cell_margins(text_cell, 0, 0, 0, 0)
    logo_p = logo_cell.paragraphs[0]
    logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_r = logo_p.add_run()
    logo_r.add_picture(str(LOGO_PATH), width=Inches(1.25))
    logo_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    text_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = text_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, "UNITED FRESH FIRE CHURCH INC.", bold=True, color=ORANGE, size=18)
    p = text_cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, "📍 Baldios, Tarlac City, Tarlac", size=10)
    p = text_cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, "☎ 09293683141", size=10)
    p = text_cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, "✉ baldemorjahaziel73@gmail.com", size=10)

    line_p = doc.add_paragraph()
    line_p.paragraph_format.space_before = Pt(4)
    line_p.paragraph_format.space_after = Pt(8)
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line_p.add_run()
    run.add_break()
    pPr = line_p._element.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), ORANGE.replace("#", ""))
    pb.append(bottom)
    pPr.append(pb)

    ref = doc.add_paragraph()
    ref.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ref.paragraph_format.space_before = Pt(0)
    ref.paragraph_format.space_after = Pt(4)
    ref.paragraph_format.line_spacing = 1.0
    add_run(ref, "Reference No.: ___________________", size=10)
    ref.add_run().add_break()
    add_run(ref, "Date: ___________________________", size=10)

    to_block = doc.add_paragraph()
    to_block.paragraph_format.space_before = Pt(2)
    to_block.paragraph_format.space_after = Pt(2)
    to_block.paragraph_format.line_spacing = 1.0
    add_run(to_block, "To:", bold=True, size=11)
    for line in [
        "The Civil Registrar General",
        "Philippine Statistics Authority",
        "PSA Complex, East Ave., Diliman",
        "Quezon City",
    ]:
        to_block.add_run().add_break()
        add_run(to_block, line, size=11)
    to_block.add_run().add_break()
    to_block.add_run().add_break()
    add_run(to_block, "Thru:", bold=True, size=11)
    for line in [
        "Atty. Sheila O. De Guzman",
        "Regional Director – RSSO I",
        "Philippine Statistics Authority",
        "City of San Fernando, La Union",
    ]:
        to_block.add_run().add_break()
        add_run(to_block, line, size=11)

    subject_label = doc.add_paragraph()
    subject_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subject_label.paragraph_format.space_before = Pt(6)
    subject_label.paragraph_format.space_after = Pt(0)
    add_run(subject_label, "Subject:", bold=True, color=ORANGE, size=11)

    subj_tbl = doc.add_table(rows=1, cols=1)
    subj_tbl.autofit = False
    subj_tbl.columns[0].width = Inches(7.0)
    remove_table_borders(subj_tbl)
    subj_cell = subj_tbl.rows[0].cells[0]
    set_cell_margins(subj_cell, 0, 0, 0, 0)
    set_cell_borders(
        subj_cell,
        bottom={"w:val": "single", "w:sz": "8", "w:space": "1", "w:color": ORANGE.replace("#", "")},
    )
    subj_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for txt in [
        "ENDORSEMENT, DESIGNATION, AND RECOMMENDATION FOR",
        "REGISTRATION OF AUTHORITY TO SOLEMNIZE MARRIAGE",
    ]:
        p = subj_cell.paragraphs[0] if not subj_cell.paragraphs[0].text else subj_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        add_run(p, txt, bold=True, color=ORANGE, size=11)

    body_lines = [
        "This is to certify that Rev. Monica Baldemor, of legal age, Filipino citizen, and residing within the Philippines, is the duly appointed Senior Pastor of United Fresh Fire Church INC., located at Baldios, Tarlac City, Tarlac.",
        "As Senior Pastor, Rev. Monica Baldemor is authorized by the Church to conduct religious services, administer church ordinances, perform ministerial duties, provide spiritual leadership to the congregation, and perform other functions pertaining to the ministry.",
        "By authority of the governing body of United Fresh Fire Church INC., we hereby designate, endorse, and recommend Rev. Monica Baldemor for the issuance of a Certificate of Registration of Authority to Solemnize Marriage (CRASM) by the Philippine Statistics Authority.",
        "We affirm that Rev. Monica Baldemor is a minister in good standing, possesses the qualifications required by our Church, and is authorized to solemnize marriages in accordance with the laws of the Republic of the Philippines and the doctrines and practices of United Fresh Fire Church INC.",
        "This certification is issued upon her request for submission to the Philippine Statistics Authority and for whatever legal purpose it may serve.",
    ]
    for text in body_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.0
        add_run(p, text, size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "Respectfully yours,", size=11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "CERTIFIED AND RECOMMENDED BY:", bold=True, color=ORANGE, size=11)

    def centered_signature(name, title, org):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.0
        add_run(p, "_____________________________", size=11)
        p.add_run().add_break()
        add_run(p, name, size=11)
        p.add_run().add_break()
        add_run(p, title, size=11)
        p.add_run().add_break()
        add_run(p, org, size=11)

    centered_signature("Noime Baldemor", "Church Secretary", "United Fresh Fire Church INC.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "ATTESTED BY:", bold=True, color=ORANGE, size=11)

    centered_signature("[Name of President / Chairman]", "President / Chairman of the Board", "United Fresh Fire Church INC.")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "WITNESSED BY:", bold=True, color=ORANGE, size=11)

    def witness_signature():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.0
        add_run(p, "_____________________________", size=11)
        p.add_run().add_break()
        add_run(p, "[Witness Name]", size=11)
        p.add_run().add_break()
        add_run(p, "[Position / Title]", size=11)
        p.add_run().add_break()
        add_run(p, "Date: ____________________", size=11)

    witness_signature()
    witness_signature()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    add_run(
        p,
        "SUBSCRIBED AND SWORN to before me this ______ day of ____________________, 20____ at ________________________________, Philippines. Affiant/s exhibited to me his/her valid identification card/s.",
        size=8.5,
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "(SEE NOTARIAL ACKNOWLEDGMENT ON THE NEXT PAGE)", italic=True, size=8)

    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    add_run(p, "NOTARIAL ACKNOWLEDGMENT", bold=True, color=ORANGE, size=12)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "REPUBLIC OF THE PHILIPPINES ) S.S.", size=11)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "CITY / MUNICIPALITY OF __________________ )", size=11)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    add_run(
        p,
        "BEFORE ME, a Notary Public for and in ____________________________, Philippines, this ______ day of ____________________, 20____ personally appeared:",
        size=10,
    )

    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [1.8, 2.25, 1.0, 1.95]
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width)
    remove_table_borders(table)
    headers = ["NAME", "VALID ID PRESENTED", "ID NO.", "DATE & PLACE ISSUED"]
    for c, text in enumerate(headers):
        cell = table.rows[0].cells[c]
        set_cell_margins(cell, 80, 80, 80, 80)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_borders(
            cell,
            top={"w:val": "single", "w:sz": "8", "w:space": "1", "w:color": "000000"},
            bottom={"w:val": "single", "w:sz": "8", "w:space": "1", "w:color": "000000"},
            left={"w:val": "single", "w:sz": "8", "w:space": "1", "w:color": "000000"},
            right={"w:val": "single", "w:sz": "8", "w:space": "1", "w:color": "000000"},
        )
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, text, bold=True, size=9)
    for r in (1, 2):
        for c in range(4):
            cell = table.rows[r].cells[c]
            set_cell_margins(cell, 80, 80, 80, 80)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(
                cell,
                top={"w:val": "single", "w:sz": "8", "w:space": "1", "w:color": "000000"},
                bottom={"w:val": "single", "w:sz": "8", "w:space": "1", "w:color": "000000"},
                left={"w:val": "single", "w:sz": "8", "w:space": "1", "w:color": "000000"},
                right={"w:val": "single", "w:sz": "8", "w:space": "1", "w:color": "000000"},
            )
            cell.paragraphs[0].paragraph_format.space_before = Pt(18)
            cell.paragraphs[0].paragraph_format.space_after = Pt(18)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    add_run(
        p,
        "known to me to be the same persons who executed the foregoing instrument and acknowledged to me that the same is their free and voluntary act and deed.",
        size=10,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    add_run(p, "WITNESS MY HAND AND SEAL this ______ day of ____________________, 20____.", size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "_____________________________", size=11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(20)
    add_run(p, "Notary Public", size=11)

    left_bottom = ["Doc. No. ______;", "Page No. ______;", "Book No. ______;", "Series of 20____."]
    right_bottom = ["PTR No. __________", "IBP No. __________", "Roll No. __________", "MCLE Compliance No. __________", "Appointment No. __________", "Until December 31, 20____"]
    footer = doc.add_table(rows=1, cols=2)
    footer.alignment = WD_TABLE_ALIGNMENT.LEFT
    footer.autofit = False
    footer.columns[0].width = Inches(3.0)
    footer.columns[1].width = Inches(4.0)
    remove_table_borders(footer)
    left, right = footer.rows[0].cells
    set_cell_margins(left, 0, 0, 0, 0)
    set_cell_margins(right, 0, 0, 0, 0)
    for cell, items, align in ((left, left_bottom, WD_ALIGN_PARAGRAPH.LEFT), (right, right_bottom, WD_ALIGN_PARAGRAPH.RIGHT)):
        for i, text in enumerate(items):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            add_run(p, text, size=9)

    return doc


def wrap_text(text: str, font_name: str, font_size: float, max_width: float):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        test = word if not current else f"{current} {word}"
        if pdfmetrics.stringWidth(test, font_name, font_size) <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_paragraph(c: canvas.Canvas, text: str, x: float, top_y: float, width: float, *, font="Helvetica", font_size=10,
                   leading=None, align=TA_LEFT, color=colors.black, bold=False, indent=0):
    font_name = "Helvetica-Bold" if bold else font
    lines = wrap_text(text, font_name, font_size, width - indent)
    if leading is None:
        leading = font_size * 1.2
    y = top_y
    for line in lines:
        if align == TA_CENTER:
            tx = x + (width - pdfmetrics.stringWidth(line, font_name, font_size)) / 2
        elif align == TA_RIGHT:
            tx = x + width - pdfmetrics.stringWidth(line, font_name, font_size)
        else:
            tx = x + indent
        c.setFont(font_name, font_size)
        c.setFillColor(color)
        c.drawString(tx, y, line)
        y -= leading
    return y


def make_pdf():
    c = canvas.Canvas(str(PDF_PATH), pagesize=legal)
    width, height = legal
    margin = 0.75 * inch
    content_w = width - 2 * margin
    orange = ORANGE_RGB

    logo_w = 1.25 * inch
    logo_h = 1.25 * inch

    def header():
        c.drawImage(str(LOGO_PATH), margin, height - margin - logo_h + 6, width=logo_w, height=logo_h, preserveAspectRatio=True, mask='auto')
        c.setFillColor(orange)
        c.setFont("Helvetica-Bold", 18)
        c.drawString(margin + logo_w + 10, height - margin - 18, "UNITED FRESH FIRE CHURCH INC.")
        c.setFillColor(colors.black)
        c.setFont("Helvetica", 10)
        c.drawString(margin + logo_w + 10, height - margin - 38, "Baldios, Tarlac City, Tarlac")
        c.drawString(margin + logo_w + 10, height - margin - 52, "09293683141")
        c.drawString(margin + logo_w + 10, height - margin - 66, "baldemorjahaziel73@gmail.com")
        c.setStrokeColor(orange)
        c.setLineWidth(1)
        c.line(margin, height - margin - 82, width - margin, height - margin - 82)

    # Page 1
    header()
    c.setFillColor(colors.black)
    c.setFont("Helvetica", 10)
    c.drawRightString(width - margin, height - margin - 104, "Reference No.: ___________________")
    c.drawRightString(width - margin, height - margin - 118, "Date: ___________________________")

    y = height - margin - 144
    y = draw_paragraph(c, "To:", margin, y, content_w, font_size=11, bold=True)
    for t in [
        "The Civil Registrar General",
        "Philippine Statistics Authority",
        "PSA Complex, East Ave., Diliman",
        "Quezon City",
    ]:
        y = draw_paragraph(c, t, margin, y - 2, content_w, font_size=11)
    y -= 6
    y = draw_paragraph(c, "Thru:", margin, y, content_w, font_size=11, bold=True)
    for t in [
        "Atty. Sheila O. De Guzman",
        "Regional Director – RSSO I",
        "Philippine Statistics Authority",
        "City of San Fernando, La Union",
    ]:
        y = draw_paragraph(c, t, margin, y - 2, content_w, font_size=11)

    c.setFillColor(orange)
    c.setFont("Helvetica-Bold", 11)
    c.drawCentredString(width / 2, y - 24, "Subject:")
    subj_y = draw_paragraph(
        c,
        "ENDORSEMENT, DESIGNATION, AND RECOMMENDATION FOR",
        margin,
        y - 40,
        content_w,
        font_size=11,
        bold=True,
        color=orange,
        align=TA_CENTER,
    )
    draw_paragraph(
        c,
        "REGISTRATION OF AUTHORITY TO SOLEMNIZE MARRIAGE",
        margin,
        subj_y - 2,
        content_w,
        font_size=11,
        bold=True,
        color=orange,
        align=TA_CENTER,
    )
    c.line(margin, subj_y - 16, width - margin, subj_y - 16)
    body_y = subj_y - 34
    body = [
        "This is to certify that Rev. Monica Baldemor, of legal age, Filipino citizen, and residing within the Philippines, is the duly appointed Senior Pastor of United Fresh Fire Church INC., located at Baldios, Tarlac City, Tarlac.",
        "As Senior Pastor, Rev. Monica Baldemor is authorized by the Church to conduct religious services, administer church ordinances, perform ministerial duties, provide spiritual leadership to the congregation, and perform other functions pertaining to the ministry.",
        "By authority of the governing body of United Fresh Fire Church INC., we hereby designate, endorse, and recommend Rev. Monica Baldemor for the issuance of a Certificate of Registration of Authority to Solemnize Marriage (CRASM) by the Philippine Statistics Authority.",
        "We affirm that Rev. Monica Baldemor is a minister in good standing, possesses the qualifications required by our Church, and is authorized to solemnize marriages in accordance with the laws of the Republic of the Philippines and the doctrines and practices of United Fresh Fire Church INC.",
        "This certification is issued upon her request for submission to the Philippine Statistics Authority and for whatever legal purpose it may serve.",
    ]
    for para in body:
        body_y = draw_paragraph(c, para, margin, body_y, content_w, font_size=10.5, leading=13, align=TA_JUSTIFY)
        body_y -= 8

    c.setFont("Helvetica", 10.5)
    c.setFillColor(colors.black)
    c.drawCentredString(width / 2, body_y - 6, "Respectfully yours,")
    c.setFillColor(orange)
    c.setFont("Helvetica-Bold", 10.5)
    c.drawCentredString(width / 2, body_y - 22, "CERTIFIED AND RECOMMENDED BY:")
    c.setFillColor(colors.black)
    for i, line in enumerate(["_____________________________", "Noime Baldemor", "Church Secretary", "United Fresh Fire Church INC."]):
        c.drawCentredString(width / 2, body_y - 44 - i * 14, line)
    c.setFillColor(orange)
    c.setFont("Helvetica-Bold", 10.5)
    c.drawCentredString(width / 2, body_y - 102, "ATTESTED BY:")
    c.setFillColor(colors.black)
    for i, line in enumerate(["_____________________________", "[Name of President / Chairman]", "President / Chairman of the Board", "United Fresh Fire Church INC."]):
        c.drawCentredString(width / 2, body_y - 124 - i * 14, line)
    c.setFillColor(orange)
    c.setFont("Helvetica-Bold", 10.5)
    c.drawCentredString(width / 2, body_y - 182, "WITNESSED BY:")
    c.setFillColor(colors.black)
    witness1_y = body_y - 204
    for i, line in enumerate(["_____________________________", "[Witness Name]", "[Position / Title]", "Date: ____________________"]):
        c.drawCentredString(width / 2, witness1_y - i * 14, line)
    witness2_y = body_y - 262
    for i, line in enumerate(["_____________________________", "[Witness Name]", "[Position / Title]", "Date: ____________________"]):
        c.drawCentredString(width / 2, witness2_y - i * 14, line)

    c.setFont("Helvetica", 8.5)
    c.drawString(margin, 76, "SUBSCRIBED AND SWORN to before me this ______ day of ____________________, 20____ at ________________________________, Philippines. Affiant/s exhibited to me his/her valid identification card/s.")
    c.setFont("Helvetica-Oblique", 8)
    c.drawString(margin, 62, "(SEE NOTARIAL ACKNOWLEDGMENT ON THE NEXT PAGE)")
    c.showPage()

    # Page 2
    c.setFillColor(orange)
    c.setFont("Helvetica-Bold", 12)
    c.drawCentredString(width / 2, height - margin - 12, "NOTARIAL ACKNOWLEDGMENT")
    c.setFillColor(colors.black)
    c.setFont("Helvetica", 10)
    c.drawString(margin, height - margin - 36, "REPUBLIC OF THE PHILIPPINES ) S.S.")
    c.drawString(margin, height - margin - 50, "CITY / MUNICIPALITY OF __________________ )")
    draw_paragraph(
        c,
        "BEFORE ME, a Notary Public for and in ____________________________, Philippines, this ______ day of ____________________, 20____ personally appeared:",
        margin,
        height - margin - 72,
        content_w,
        font_size=9.5,
    )

    table_x = margin
    table_y = height - margin - 126
    row_h = 28
    col_w = [1.85 * inch, 2.25 * inch, 1.0 * inch, 1.9 * inch]
    headers = ["NAME", "VALID ID PRESENTED", "ID NO.", "DATE & PLACE ISSUED"]
    table_data = [headers, ["", "", "", ""], ["", "", "", ""]]
    tbl = Table(table_data, colWidths=col_w, rowHeights=[row_h, 1.0 * inch, 0.95 * inch])
    tbl.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 9),
        ("ALIGN", (0, 0), (-1, 0), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.8, colors.black),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 1), (-1, -1), 8.5),
        ("LEADING", (0, 0), (-1, -1), 10),
    ]))
    w, h = tbl.wrapOn(c, content_w, 200)
    tbl.drawOn(c, table_x, table_y - h + 10)

    draw_paragraph(
        c,
        "known to me to be the same persons who executed the foregoing instrument and acknowledged to me that the same is their free and voluntary act and deed.",
        margin,
        table_y - h - 18,
        content_w,
        font_size=9.5,
    )
    c.drawString(margin, table_y - h - 48, "WITNESS MY HAND AND SEAL this ______ day of ____________________, 20____.")
    c.drawRightString(width - margin, table_y - h - 50, "_____________________________")
    c.drawRightString(width - margin, table_y - h - 64, "Notary Public")

    c.setFont("Helvetica", 9)
    for i, text in enumerate(["Doc. No. ______;", "Page No. ______;", "Book No. ______;", "Series of 20____."]):
        c.drawString(margin, 72 - i * 14, text)
    right_lines = [
        "PTR No. __________",
        "IBP No. __________",
        "Roll No. __________",
        "MCLE Compliance No. __________",
        "Appointment No. __________",
        "Until December 31, 20____",
    ]
    for i, text in enumerate(right_lines):
        c.drawRightString(width - margin, 72 - i * 14, text)

    c.save()


def main():
    ensure_dir(OUT_DIR)
    doc = make_docx()
    doc.save(DOCX_PATH)
    make_pdf()
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote {PDF_PATH}")


if __name__ == "__main__":
    main()
